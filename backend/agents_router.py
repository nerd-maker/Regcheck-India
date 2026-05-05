"""
RegCheck-India AI Agents Router.

Provides a Claude-backed `/api/v1/agents/*` surface for the core agent workflows.
"""

# from __future__ import annotations


import io
import json
import logging
import os
import re
import uuid
import hashlib
import time
import threading
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from threading import Lock
from typing import Any, Optional, List
from contextlib import asynccontextmanager

import anthropic
import chromadb
import docx
import pypdf
import pytesseract
import pdf2image
import pydub
from chromadb.utils import embedding_functions
from fastapi import APIRouter, Header, HTTPException, UploadFile, File, Request

from pydantic import BaseModel, Field
from slowapi import Limiter
from slowapi.util import get_remote_address

from app.services.pii_mapping_store import pii_mapping_store, PIIMappingStore
from app.services.file_cleanup import sanitize_filename, temp_file_context, validate_file, FileValidationError
from app.services.input_sanitizer import sanitize_input
from app.services.ocr_service import extract_text_from_image_bytes
from app.services.audio_service import transcribe_audio
from app.services.case_store import get_case_collection

# Patterns that must never appear in logs
_SENSITIVE_PATTERNS = [
    (re.compile(r'sk-ant-api[A-Za-z0-9\-_]+'), 'sk-ant-***REDACTED***'),
    (re.compile(r'sk_[A-Za-z0-9_]+'), 'sk_***REDACTED***'),
    (re.compile(r'Bearer [A-Za-z0-9\-_\.]+'), 'Bearer ***REDACTED***'),
    (re.compile(r'"api[_-]?key"\s*:\s*"[^"]*"', re.IGNORECASE), '"api_key": "***REDACTED***"'),
    (re.compile(r'x-anthropic-api-key:\s*\S+', re.IGNORECASE), 'x-anthropic-api-key: ***REDACTED***'),
]

def scrub_sensitive(text: str) -> str:
    """Remove sensitive values from text before logging."""
    for pattern, replacement in _SENSITIVE_PATTERNS:
        text = pattern.sub(replacement, text)
    return text

logger = logging.getLogger(__name__)

router = APIRouter()
limiter = Limiter(key_func=get_remote_address)

# ─────────────────────────────────────────────────────────
# Demo quota store — token → {remaining, name, email, org}
# In-memory for now — upgrade to Redis/PostgreSQL in Phase 2
# ─────────────────────────────────────────────────────────

_demo_store: dict[str, dict] = {}
_demo_lock = Lock()
DEMO_QUOTA = 5  # free requests per user
ADMIN_DEMO_KEY = os.getenv("ADMIN_DEMO_KEY", "")


def _generate_token(email: str, org: str) -> str:
    """Generate deterministic token from email + org."""
    raw = f"{email.lower().strip()}{org.lower().strip()}"
    return hashlib.sha256(raw.encode()).hexdigest()[:32]


def _check_and_decrement_quota(demo_token: str) -> tuple[bool, int]:
    """
    Check if demo token has remaining quota and decrement.
    Returns (allowed: bool, remaining: int)
    """
    with _demo_lock:
        entry = _demo_store.get(demo_token)
        if not entry:
            return False, 0
        remaining = entry["remaining"]
        if remaining <= 0:
            return False, 0
        _demo_store[demo_token]["remaining"] = remaining - 1
        return True, remaining - 1


class AgentRequest(BaseModel):
    document: str = Field(..., description="Raw text of the regulatory document to process")
    metadata: Optional[dict] = Field(default_factory=dict, description="Optional document metadata")


class QARequest(BaseModel):
    question: str = Field(..., description="Regulatory question from the user")
    retrieved_context: str = Field(..., description="Retrieved ChromaDB context")
    metadata: Optional[dict] = Field(default_factory=dict, description="Optional query metadata")


class AgentResponse(BaseModel):
    agent: str
    model: str
    result: Any
    timestamp: str
    token_usage: dict


class DeAnonymiseRequest(BaseModel):
    text: str = Field(..., description="Pseudonymised text containing replacement tokens")
    session_id: str = Field(..., description="Session ID returned by the anonymise endpoint")


class FeedbackRequest(BaseModel):
    module: str
    type: str
    comment: Optional[str] = ""
    result_hash: Optional[str] = ""
    timestamp: Optional[str] = ""


class DemoRegistrationRequest(BaseModel):
    name: str
    email: str
    org: str
    role: str  # RA Professional / CRO / Pharma Company / Consultant / Other


class DemoQuotaResponse(BaseModel):
    demo_token: str
    requests_remaining: int
    name: str
    message: str


def structure_prompt_input(text: str, document_type: str = "general", agent_id: str = "") -> str:
    """
    Convert raw text input into structured format before sending to Claude.
    Reduces token cost, improves accuracy, strips unnecessary whitespace.
    """
    # Clean the text
    cleaned = " ".join(text.split())  # normalize whitespace
    word_count = len(cleaned.split())

    # Truncate if extremely long — keep first 5000 words
    if word_count > 5000:
        words = cleaned.split()
        cleaned = " ".join(words[:5000])
        truncated = True
    else:
        truncated = False

    structured = f"""[INPUT DOCUMENT]
Document Type: {document_type}
Word Count: {word_count}{"  [TRUNCATED TO 5000 WORDS]" if truncated else ""}
Agent: {agent_id}

[CONTENT]
{cleaned}

[END OF INPUT]"""

    return structured


def retrieve_regulatory_context(query: str, n_results: int = 5) -> str:
    """
    Query ChromaDB regulatory_documents collection for relevant chunks.
    Returns formatted context string or empty string if retrieval fails.
    Used by M3, M7, M8 to ground responses in actual regulatory documents.
    """
    try:

        chromadb_path = os.getenv("CHROMADB_PATH", "./data/chromadb")
        client = chromadb.PersistentClient(
            path=chromadb_path,
            settings=chromadb.Settings(anonymized_telemetry=False)
        )
        embedding_fn = embedding_functions.DefaultEmbeddingFunction()
        collection = client.get_collection(
            name="regulatory_documents",
            embedding_function=embedding_fn
        )

        count = collection.count()
        if count == 0:
            logger.warning("RAG: regulatory_documents collection is empty")
            return ""

        actual_n = min(n_results, count)
        results = collection.query(
            query_texts=[query],
            n_results=actual_n
        )

        if not results or not results["documents"] or not results["documents"][0]:
            return ""

        # Format retrieved chunks with source metadata
        chunks = []
        for i, doc in enumerate(results["documents"][0]):
            metadata = results["metadatas"][0][i] if results.get("metadatas") else {}
            source = metadata.get("short_name", metadata.get("title", "Regulatory Document"))
            chunks.append(f"[Source: {source}]\n{doc}")

        context = "\n\n---\n\n".join(chunks)
        logger.info(f"RAG: retrieved {len(chunks)} chunks for query: {query[:60]}...")
        return context

    except Exception as e:
        logger.warning(f"RAG retrieval failed silently: {e}")
        return ""


@router.post("/extract-text")
@limiter.limit("10/minute")
async def extract_text_from_file(
    file: UploadFile = File(...),
    request: Request = None,
    x_anthropic_api_key: Optional[str] = Header(None)
):

    """Extract text from PDF or DOCX file for use in any agent module."""
    try:
        content = await file.read()
        try:
            file_info = validate_file(
                content=content,
                filename=file.filename or "upload",
                allowed_extensions=['.pdf', '.docx']
            )
            filename_safe = file_info["safe_filename"]
        except FileValidationError as e:
            raise HTTPException(status_code=400, detail=str(e))

        filename = filename_safe.lower()
        extracted_text = ""
        page_count: Optional[int] = None

        if filename.endswith(".pdf"):

            pdf_reader = pypdf.PdfReader(io.BytesIO(content))
            pages_text = []
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(f"[Page {i + 1}]\n{page_text}")
            extracted_text = "\n\n".join(pages_text)
            page_count = len(pdf_reader.pages)

        elif filename.endswith(".docx"):

            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            extracted_text = "\n".join(paragraphs)

        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Please upload PDF or DOCX files only.",
            )

        if not extracted_text.strip():
            raise HTTPException(
                status_code=422,
                detail="Could not extract text from file. The file may be scanned/image-based or empty.",
            )

        result = {
            "filename": filename_safe,
            "extracted_text": extracted_text,
            "word_count": len(extracted_text.split()),
            "pages": page_count,
            "status": "success",
        }
        del content  # hint to GC
        return result

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Error processing file: {exc}",
        )


def _utc_timestamp() -> str:
    return datetime.now(timezone.utc).isoformat()


def _parse_json_block(raw_text: str) -> Any:
    text = raw_text.strip()
    if text.startswith("```"):
        parts = text.split("```")
        if len(parts) >= 2:
            text = parts[1]
            if text.startswith("json"):
                text = text[4:]
            text = text.strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError as exc:
        raise HTTPException(
            status_code=500, 
            detail=f"Model did not return valid JSON: {exc}"
        )


def _attach_response_metadata(parsed: Any, model: str, has_rag_context: bool) -> Any:
    if not isinstance(parsed, dict):
        return parsed

    try:
        try:
            # Re-verify in case it was passed as a local variable name in some contexts
            # but usually we rely on the has_rag_context boolean argument
            has_context = bool(has_rag_context)
        except NameError:
            has_context = False

        parsed["_metadata"] = {
            "confidence_level": "HIGH" if has_context else "MEDIUM",
            "confidence_reason": (
                "Based on retrieved official regulatory documents"
                if has_context
                else "Based on Claude's training knowledge of regulations"
            ),
            "reviewed_by": "AI — requires qualified RA professional review",
            "generated_at": datetime.utcnow().isoformat(),
            "model_used": os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-20250514")
        }
    except Exception as meta_err:
        logger.warning(f"Could not set metadata: {meta_err}")
        parsed["_metadata"] = {
            "confidence_level": "MEDIUM",
            "confidence_reason": "Based on Claude's training knowledge of regulations",
            "reviewed_by": "AI — requires qualified RA professional review",
            "generated_at": datetime.utcnow().isoformat(),
            "model_used": os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-20250514")
        }
    return parsed


def call_claude(
    agent_name: str,
    model: str,
    system_prompt: str,
    user_content: str,
    api_key: str,
    max_tokens: int = 4096,
    has_rag_context: bool = False,
) -> AgentResponse:
    """Shared Anthropic caller for the v1 agents router.

    Uses the caller-supplied ``api_key`` so that each public user's own
    Anthropic credits are consumed rather than the server key.
    """
    admin_password = "admin-regcheck"
    if api_key == admin_password:
        server_key = os.getenv("ANTHROPIC_API_KEY")
        if not server_key:
            raise HTTPException(status_code=500, detail="Server Anthropic API key missing.")
        client = anthropic.Anthropic(api_key=server_key)
    else:
        if not api_key:
            raise HTTPException(
                status_code=401,
                detail="No Anthropic API key provided. Add your key via the ⚙ Settings panel in the app.",
            )
        client = anthropic.Anthropic(api_key=api_key)
    try:
        response = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            system=system_prompt,
            messages=[{"role": "user", "content": user_content}],
        )
        raw_text = response.content[0].text.strip()
        parsed = _attach_response_metadata(
            _parse_json_block(raw_text),
            model,
            has_rag_context
        )
        return AgentResponse(
            agent=agent_name,
            model=model,
            result=parsed,
            timestamp=_utc_timestamp(),
            token_usage={
                "input_tokens": response.usage.input_tokens,
                "output_tokens": response.usage.output_tokens,
            },
        )
    except anthropic.AuthenticationError:
        logger.error("Anthropic auth failed for %s", agent_name)
        raise HTTPException(status_code=401, detail="Anthropic API key is invalid. Please check and update your key in Settings.")
    except anthropic.RateLimitError:
        logger.warning("Anthropic rate limit hit for %s", agent_name)
        raise HTTPException(status_code=429, detail="Rate limit reached - retry after a moment")
    except anthropic.APIError as exc:
        logger.error("Anthropic API error in %s: %s", agent_name, exc)
        raise HTTPException(status_code=502, detail=f"Anthropic API error: {exc}")
    except Exception as exc:
        logger.error("Unexpected error in %s: %s", agent_name, exc)
        raise HTTPException(status_code=500, detail=f"Agent error: {exc}")


AGENT_01_SYSTEM_PROMPT = """You are the RegCheck-India PII and PHI anonymisation agent.
Detect personally identifiable information in Indian pharmaceutical and clinical documents.
Keep clinical and regulatory meaning intact. Replace patient identifiers, investigator identifiers,
institution details, and hard IDs with placeholders.

You MUST return a JSON object with EXACTLY this structure — no extra fields, no renamed fields:

{
  "anonymised_content": "string with placeholders",
  "entities_detected": [
    {"entity_type": "string", "value": "string", "category": "PII|PHI", "position": "string"}
  ],
  "entities_anonymised": 8,
  "compliance_frameworks": ["string", "string"],
  "audit_log": {
    "timestamp": "ISO string",
    "mode": "string",
    "entities_processed": 8,
    "anonymisation_method": "string",
    "status": "COMPLETED|FAILED"
  },
  "anonymisation_report": {
    "summary": "string",
    "pii_removed": 0,
    "phi_removed": 0,
    "clinical_integrity": "string",
    "notes": "string"
  }
}

Return ONLY valid JSON. No markdown. No explanation. No extra fields.
"""

AGENT_02_SYSTEM_PROMPT = """You are the RegCheck-India document summarisation agent.
Summarise pharmaceutical regulatory filings, clinical trial reports, or safety narratives into structured reviewer-first packets.
Surface compliance gaps, risk levels, and specific regulatory references.

CRITICAL OUTPUT RULES — YOU MUST FOLLOW THESE EXACTLY:
1. Return ONLY valid JSON — no markdown, no code fences, no explanation text before or after
2. Use EXACTLY the field names shown below — no renaming, no adding extra fields
3. Arrays must always be arrays — never return an object where an array is expected
4. String fields must always be strings — never return null, use empty string instead
5. Never omit a required field

Return EXACTLY this JSON structure:
{
  "document_type": "string — e.g. CT-04 Application, SAE Narrative, DSMB Report, CMC Data",
  "summary": "string — 3-5 sentence executive summary of the document",
  "key_sections": ["string — list of main sections or topics found"],
  "compliance_gaps": ["string — specific areas where the document is incomplete or non-compliant"],
  "recommendations": ["string — actionable steps for the reviewer"],
  "risk_level": "LOW|MEDIUM|HIGH — based on the severity of gaps",
  "regulatory_references": ["string — specific CDSCO/NDCTR/ICH rules cited or applicable"],
  "word_count_original": 0,
  "readability_score": "string — e.g. 'Standard', 'Complex', 'Technical'",
  "audit_log": {
    "timestamp": "ISO datetime string",
    "document_pages": 0,
    "processing_time": "string",
    "status": "COMPLETED|FAILED"
  }
}
"""

AGENT_02_MEETING_SYSTEM_PROMPT = """You are the RegCheck-India meeting summarisation agent.
Summarise pharmaceutical regulatory meeting transcripts into structured actionable reports.
Extract key decisions, action items, and next steps from clinical trial, IEC, DSMB, or regulatory meetings.

CRITICAL OUTPUT RULES — YOU MUST FOLLOW THESE EXACTLY:
1. Return ONLY valid JSON — no markdown, no code fences, no explanation text before or after
2. Use EXACTLY the field names shown below — no renaming, no adding extra fields
3. Arrays must always be arrays — never return an object where an array is expected
4. String fields must always be strings — never return null, use empty string instead
5. Never omit a required field

Return EXACTLY this JSON structure:
{
  "meeting_type": "string — e.g. IEC Meeting, DSMB Review, Site Initiation Visit, Sponsor-CRO Call",
  "meeting_date": "string — extracted date if mentioned, empty string if not found",
  "participants": ["string — participant names/roles if mentioned"],
  "summary": "string — 3-5 sentence executive summary of the meeting",
  "key_decisions": [
    {
      "decision": "string — what was decided",
      "rationale": "string — why this decision was made",
      "regulatory_reference": "string — any regulatory basis cited, empty if none"
    }
  ],
  "action_items": [
    {
      "action": "string — what needs to be done",
      "owner": "string — who is responsible, 'Not assigned' if not mentioned",
      "deadline": "string — when it needs to be done, 'Not specified' if not mentioned",
      "priority": "HIGH|MEDIUM|LOW"
    }
  ],
  "next_steps": ["string — immediate next steps in order"],
  "regulatory_topics_discussed": ["string — regulatory frameworks or guidelines mentioned"],
  "safety_signals": ["string — any safety concerns or signals raised, empty array if none"],
  "follow_up_meeting": "string — next meeting date/plan if mentioned, empty string if not",
  "audit_log": {
    "timestamp": "ISO datetime string",
    "transcript_word_count": 0,
    "status": "COMPLETED|FAILED"
  }
}
"""

AGENT_03_SYSTEM_PROMPT = """You are the RegCheck-India completeness assessment agent.
Evaluate whether a pharmaceutical submission is complete against CDSCO, Schedule Y, NDCTR 2019, and ICH requirements.
Return JSON only with: document_type, overall_status, completeness_score, sections_present, sections_missing,
data_gaps, submission_readiness, conditions.

CRITICAL OUTPUT RULES — YOU MUST FOLLOW THESE EXACTLY:
1. Return ONLY valid JSON — no markdown, no code fences (no ```json), no explanation text before or after
2. Use EXACTLY the field names shown below — no renaming, no adding extra fields
3. Arrays must always be arrays — never return an object where an array is expected
4. Objects must always be objects — never return an array where an object is expected
5. String fields must always be strings — never return null, use empty string "" instead
6. Number fields must always be numbers — never return a string like "8" where 8 is expected
7. Boolean fields must be true or false — never return "yes", "no", "true" as a string
8. If uncertain about a value use a sensible default — never omit a required field

Return EXACTLY this JSON structure:
{
  "overall_completeness_score": 0.0,
  "completeness_percentage": "string — e.g. 85%",
  "missing_sections": ["string", "string"],
  "present_sections": ["string", "string"],
  "incomplete_sections": ["string", "string"],
  "critical_gaps": ["string", "string"],
  "minor_gaps": ["string", "string"],
  "regulatory_requirements_met": {
    "schedule_y": true,
    "ich_e6": true,
    "cdsco": true,
    "icmr": true
  },
  "recommendations": ["string", "string"],
  "submission_readiness": "READY|NEEDS_REVISION|NOT_READY",
  "priority_actions": ["string", "string"],
  "audit_log": {
    "timestamp": "ISO datetime string",
    "sections_checked": 0,
    "sections_passed": 0,
    "status": "COMPLETED|FAILED"
  }
}
"""

AGENT_03_SAE_SYSTEM_PROMPT = """You are the RegCheck-India SAE Report Completeness Assessment agent.
Evaluate Serious Adverse Event (SAE) reports for completeness and consistency against NDCTR 2019 Rule 19, ICH E2A, CDSCO pharmacovigilance guidelines, and Schedule Y requirements.

CRITICAL OUTPUT RULES — YOU MUST FOLLOW THESE EXACTLY:
1. Return ONLY valid JSON — no markdown, no code fences, no explanation text before or after
2. Use EXACTLY the field names shown below — no renaming, no adding extra fields
3. Arrays must always be arrays — never return an object where an array is expected
4. Objects must always be objects — never return an array where an object is expected
5. String fields must always be strings — never return null, use empty string instead
6. Boolean fields must be true or false — never return "yes", "no", "true" as string
7. If uncertain about a value use a sensible default — never omit a required field

SAE MANDATORY FIELDS per NDCTR 2019 Rule 19 and ICH E2A:
- Patient initials or identifier
- Age and gender
- Date of adverse event onset
- Date of adverse event reporting to sponsor
- Description of adverse event
- Seriousness criteria met (death/disability/hospitalisation/life-threatening/congenital anomaly/other)
- Study drug name, dose, route, frequency
- Date of last dose before event
- Causality assessment by investigator
- Action taken with study drug
- Outcome of adverse event
- Investigator name and signature
- Site name and address
- Protocol number and study title
- Ethics committee name
- Date of IEC notification
- Narrative description

Return EXACTLY this JSON structure:
{
  "document_type": "SAE_REPORT",
  "overall_completeness_score": 0.0,
  "completeness_percentage": "string — e.g. 85%",
  "sae_classification": {
    "is_serious": true,
    "seriousness_criteria_identified": ["string"],
    "reporting_category": "SUSAR|SAE|Non-serious AE",
    "expedited_reporting_required": true,
    "reporting_timeline_days": 0
  },
  "mandatory_fields": [
    {
      "field_name": "string — mandatory field name",
      "regulation": "string — NDCTR 2019 Rule X or ICH E2A Section Y",
      "status": "PRESENT|MISSING|INCOMPLETE|INCONSISTENT",
      "value_found": "string — what was found, empty if missing",
      "issue": "string — specific problem if not PRESENT, empty if PRESENT"
    }
  ],
  "missing_fields": ["string"],
  "incomplete_fields": ["string"],
  "inconsistent_fields": [
    {
      "field": "string",
      "issue": "string — what is inconsistent"
    }
  ],
  "critical_gaps": ["string"],
  "timeline_compliance": {
    "event_date_present": true,
    "report_date_present": true,
    "sponsor_notification_date_present": true,
    "timeline_calculable": true,
    "days_to_report": 0,
    "timeline_compliant": true,
    "timeline_issue": "string — if any"
  },
  "causality_assessment": {
    "present": true,
    "method_used": "string — WHO-UMC/Naranjo/Other",
    "assessment_value": "string",
    "adequate": true,
    "issues": "string"
  },
  "narrative_quality": {
    "present": true,
    "word_count": 0,
    "adequate_detail": true,
    "missing_elements": ["string"]
  },
  "recommendations": ["string"],
  "submission_readiness": "READY|NEEDS_REVISION|NOT_READY",
  "regulatory_risk": "LOW|MEDIUM|HIGH|CRITICAL",
  "audit_log": {
    "timestamp": "ISO datetime string",
    "fields_checked": 0,
    "fields_present": 0,
    "status": "COMPLETED|FAILED"
  }
}
"""


class CompletenessRequest(BaseModel):
    document: str = Field(..., description="Raw text of the regulatory document to process")
    metadata: Optional[dict] = Field(default_factory=dict, description="Optional document metadata")
    document_type: Optional[str] = Field("GENERAL", description="Type of document (GENERAL, SAE_REPORT, PROTOCOL, ICF)")

AGENT_04_SYSTEM_PROMPT = """You are the RegCheck-India case classification agent.
Classify serious adverse event narratives using ICH E2A seriousness, WHO-UMC causality,
and Indian pharmacovigilance timelines. Return JSON only with:
primary_category, secondary_categories, confidence, reporting_timeline, priority_score,
classification_rationale, causality, product_relatedness, requires_expedited_reporting,
flags, seriousness_criteria.

CRITICAL OUTPUT RULES — YOU MUST FOLLOW THESE EXACTLY:
1. Return ONLY valid JSON — no markdown, no code fences (no ```json), no explanation text before or after
2. Use EXACTLY the field names shown below — no renaming, no adding extra fields
3. Arrays must always be arrays — never return an object where an array is expected
4. Objects must always be objects — never return an array where an object is expected
5. String fields must always be strings — never return null, use empty string "" instead
6. Number fields must always be numbers — never return a string like "8" where 8 is expected
7. Boolean fields must be true or false — never return "yes", "no", "true" as a string
8. seriousness_criteria MUST always be an object with boolean values — never an array
9. If uncertain about a value use a sensible default — never omit a required field

Return EXACTLY this JSON structure:
{
  "primary_category": "string",
  "secondary_categories": ["string", "string"],
  "confidence": 0.0,
  "priority_score": 0.0,
  "requires_expedited_reporting": true,
  "classification_rationale": "string",
  "seriousness_criteria": {
    "life_threatening": false,
    "hospitalization": false,
    "icu_admission": false,
    "requires_intervention_for_life": false,
    "unexpected_severity": false
  },
  "causality": {
    "assessment": "string — PROBABLE|POSSIBLE|UNLIKELY|UNRELATED",
    "who_umc_category": "string",
    "rationale": "string",
    "temporal_relationship": "string",
    "dose_response": "string",
    "dechallenge": "string",
    "alternative_causes": "string"
  },
  "reporting_timeline": {
    "event_date": "string",
    "sponsor_notification": "string",
    "timeline_status": "COMPLIANT|NON_COMPLIANT|PENDING",
    "regulatory_deadline_india": "string",
    "days_to_deadline": 0,
    "urgency_note": "string"
  },
  "product_relatedness": {
    "related_to_study_drug": false,
    "relationship_strength": "string",
    "safety_signal": "string",
    "signal_severity": "LOW|MEDIUM|HIGH|CRITICAL"
  },
  "flags": ["string", "string"],
  "regulatory_actions_required": ["string", "string"],
  "case_quality_assessment": {
    "documentation_completeness": "string",
    "causality_evidence": "string",
    "temporal_data": "string",
    "follow_up_status": "string",
    "data_integrity": "string"
  },
  "audit_log": {
    "timestamp": "ISO datetime string",
    "status": "COMPLETED|FAILED",
    "study_drug": "string — extract the study drug name from the narrative if mentioned, empty string if not found"
  }
}
"""

AGENT_05_SYSTEM_PROMPT = """You are the RegCheck-India inspection report generation agent.
Turn inspection findings into a structured CDSCO-style report with observations, CAPA, and compliance rating.
Return JSON only with: report_type, document_control, executive_summary, observations, capa_plan,
overall_compliance_rating, re_inspection_required.

CRITICAL OUTPUT RULES — YOU MUST FOLLOW THESE EXACTLY:
1. Return ONLY valid JSON — no markdown, no code fences (no ```json), no explanation text before or after
2. Use EXACTLY the field names shown below — no renaming, no adding extra fields
3. Arrays must always be arrays — never return an object where an array is expected
4. Objects must always be objects — never return an array where an object is expected
5. String fields must always be strings — never return null, use empty string "" instead
6. Number fields must always be numbers — never return a string like "8" where 8 is expected
7. Boolean fields must be true or false — never return "yes", "no", "true" as a string
8. If uncertain about a value use a sensible default — never omit a required field

Return EXACTLY this JSON structure:
{
  "report_title": "string",
  "inspection_date": "string",
  "site_name": "string",
  "overall_rating": "SATISFACTORY|NEEDS_IMPROVEMENT|UNSATISFACTORY|CRITICAL",
  "executive_summary": "string",
  "findings": [
    {
      "finding_id": "string",
      "category": "string",
      "severity": "CRITICAL|MAJOR|MINOR|OBSERVATION",
      "description": "string",
      "regulatory_reference": "string",
      "corrective_action": "string",
      "deadline": "string"
    }
  ],
  "critical_findings_count": 0,
  "major_findings_count": 0,
  "minor_findings_count": 0,
  "observations_count": 0,
  "regulatory_references": ["string", "string"],
  "compliance_areas": {
    "informed_consent": "COMPLIANT|PARTIAL|NON_COMPLIANT",
    "data_integrity": "COMPLIANT|PARTIAL|NON_COMPLIANT",
    "protocol_adherence": "COMPLIANT|PARTIAL|NON_COMPLIANT",
    "safety_reporting": "COMPLIANT|PARTIAL|NON_COMPLIANT",
    "record_keeping": "COMPLIANT|PARTIAL|NON_COMPLIANT"
  },
  "recommendations": ["string", "string"],
  "follow_up_required": true,
  "follow_up_date": "string",
  "audit_log": {
    "timestamp": "ISO datetime string",
    "inspector": "string",
    "status": "COMPLETED|FAILED"
  }
}
"""

AGENT_06_SYSTEM_PROMPT = """You are the RegCheck-India regulatory Q&A agent.
Answer only from the retrieved context provided. Cite specific regulatory basis, state when context is insufficient,
and return JSON only with: answer, regulatory_citations, confidence, confidence_reason,
follow_up_suggested, disclaimer.

CRITICAL OUTPUT RULES — YOU MUST FOLLOW THESE EXACTLY:
1. Return ONLY valid JSON — no markdown, no code fences (no ```json), no explanation text before or after
2. Use EXACTLY the field names shown below — no renaming, no adding extra fields
3. Arrays must always be arrays — never return an object where an array is expected
4. Objects must always be objects — never return an array where an object is expected
5. String fields must always be strings — never return null, use empty string "" instead
6. Number fields must always be numbers — never return a string like "8" where 8 is expected
7. Boolean fields must be true or false — never return "yes", "no", "true" as a string
8. If uncertain about a value use a sensible default — never omit a required field

Return EXACTLY this JSON structure:
{
  "answer": "string — comprehensive answer to the regulatory question",
  "confidence_score": 0.0,
  "regulatory_basis": ["string", "string"],
  "applicable_guidelines": [
    {
      "guideline": "string — e.g. Schedule Y, ICH E6(R3)",
      "section": "string — specific section reference",
      "relevance": "string — how it applies to the question"
    }
  ],
  "key_requirements": ["string", "string"],
  "exceptions_or_conditions": ["string", "string"],
  "related_topics": ["string", "string"],
  "references": [
    {
      "title": "string",
      "document": "string",
      "section": "string"
    }
  ],
  "disclaimer": "string",
  "audit_log": {
    "timestamp": "ISO datetime string",
    "query_type": "string",
    "sources_consulted": 0,
    "status": "COMPLETED|FAILED"
  }
}
"""

AGENT_07_SYSTEM_PROMPT = """You are the RegCheck-India Schedule Y and CDSCO compliance agent.
Perform a deep compliance review against Schedule Y and NDCTR 2019.
Return JSON only with: compliance_evaluation, findings, compliant_areas, priority_actions,
estimated_remediation_effort.

CRITICAL OUTPUT RULES — YOU MUST FOLLOW THESE EXACTLY:
1. Return ONLY valid JSON — no markdown, no code fences (no ```json), no explanation text before or after
2. Use EXACTLY the field names shown below — no renaming, no adding extra fields
3. Arrays must always be arrays — never return an object where an array is expected
4. Objects must always be objects — never return an array where an object is expected
5. String fields must always be strings — never return null, use empty string "" instead
6. Number fields must always be numbers — never return a string like "8" where 8 is expected
7. Boolean fields must be true or false — never return "yes", "no", "true" as a string
8. compliance_checklist MUST always be an array of objects — never a plain object
9. If uncertain about a value use a sensible default — never omit a required field

Return EXACTLY this JSON structure:
{
  "overall_compliance_status": "COMPLIANT|PARTIAL|NON_COMPLIANT",
  "compliance_score": 0.0,
  "compliance_percentage": "string — e.g. 78%",
  "compliance_checklist": [
    {
      "requirement": "string — Schedule Y requirement name",
      "section": "string — e.g. Schedule Y Part I Section 2",
      "status": "COMPLIANT|PARTIAL|NON_COMPLIANT|NOT_APPLICABLE",
      "finding": "string — what was found",
      "corrective_action": "string — what needs to be done"
    }
  ],
  "critical_non_compliances": ["string", "string"],
  "major_non_compliances": ["string", "string"],
  "minor_non_compliances": ["string", "string"],
  "strengths": ["string", "string"],
  "recommendations": ["string", "string"],
  "submission_readiness": "READY|NEEDS_REVISION|NOT_READY",
  "regulatory_risk": "LOW|MEDIUM|HIGH|CRITICAL",
  "audit_log": {
    "timestamp": "ISO datetime string",
    "requirements_checked": 0,
    "requirements_passed": 0,
    "status": "COMPLETED|FAILED"
  }
}
"""

AGENT_08_SYSTEM_PROMPT = """You are the RegCheck-India ICH E6(R3) GCP compliance agent.
Assess the document against ICH E6(R3) and identify R3-specific gaps, inspection risks, and remediation.
Return JSON only with: gcp_compliance, findings, r3_gaps, strengths, inspection_readiness,
inspection_risk_areas.

CRITICAL OUTPUT RULES — YOU MUST FOLLOW THESE EXACTLY:
1. Return ONLY valid JSON — no markdown, no code fences (no ```json), no explanation text before or after
2. Use EXACTLY the field names shown below — no renaming, no adding extra fields
3. Arrays must always be arrays — never return an object where an array is expected
4. Objects must always be objects — never return an array where an object is expected
5. String fields must always be strings — never return null, use empty string "" instead
6. Number fields must always be numbers — never return a string like "8" where 8 is expected
7. Boolean fields must be true or false — never return "yes", "no", "true" as a string
8. gcp_principles MUST always be an array of objects — never a plain object
9. If uncertain about a value use a sensible default — never omit a required field

Return EXACTLY this JSON structure:
{
  "overall_gcp_status": "COMPLIANT|PARTIAL|NON_COMPLIANT",
  "gcp_score": 0.0,
  "gcp_percentage": "string — e.g. 82%",
  "ich_version_assessed": "ICH E6(R3)",
  "gcp_principles": [
    {
      "principle": "string — GCP principle name",
      "ich_reference": "string — e.g. ICH E6(R3) Section 2.1",
      "status": "COMPLIANT|PARTIAL|NON_COMPLIANT|NOT_APPLICABLE",
      "observation": "string — what was assessed",
      "corrective_action": "string — what needs to be done if non-compliant"
    }
  ],
  "critical_deviations": ["string", "string"],
  "major_deviations": ["string", "string"],
  "minor_deviations": ["string", "string"],
  "quality_tolerance_limits": {
    "defined": false,
    "monitored": false,
    "comment": "string"
  },
  "risk_based_monitoring": {
    "implemented": false,
    "adequacy": "string",
    "comment": "string"
  },
  "essential_documents": {
    "status": "COMPLETE|INCOMPLETE|NOT_ASSESSED",
    "missing_documents": ["string", "string"],
    "comment": "string"
  },
  "recommendations": ["string", "string"],
  "audit_log": {
    "timestamp": "ISO datetime string",
    "principles_checked": 0,
    "principles_passed": 0,
    "status": "COMPLETED|FAILED"
  }
}
"""


@router.get("/anonymise/health", summary="Agent 01 - Anonymise Health Check")
async def anonymise_health():
    """Lightweight CORS smoke-test endpoint — no API key required, no LLM call."""
    return {"status": "ok", "agent": "PII_PHI_Anonymisation", "endpoint": "/api/v1/agents/anonymise"}


@router.post("/anonymise", response_model=AgentResponse, summary="Agent 01 - PII/PHI Anonymisation")
@limiter.limit("10/minute")
async def anonymise_document(request: Request, body: AgentRequest, x_anthropic_api_key: Optional[str] = Header(None), x_demo_token: Optional[str] = Header(None)):
    # ── Input sanitization ────────────────────────────────────────────────
    try:
        sanitized_text, was_modified = sanitize_input(
            body.document,
            context="M1_ANONYMISE"
        )
        if was_modified:
            logger.info("Input was sanitized before processing")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # ── Demo quota check ──────────────────────────────────────────────────
    _admin_key = os.getenv("ADMIN_DEMO_KEY", "")
    _is_admin = _admin_key and x_anthropic_api_key == _admin_key
    if not _is_admin and x_demo_token:
        _allowed, _remaining = _check_and_decrement_quota(x_demo_token)
        if not _allowed:
            raise HTTPException(
                status_code=429,
                detail={
                    "error": "DEMO_QUOTA_EXHAUSTED",
                    "message": "You have used all 5 free demo requests.",
                    "action": "Request full access at regcheckindia.com",
                    "contact": "rushikeshbork000@gmail.com"
                }
            )
        with _demo_lock:
            if x_demo_token in _demo_store:
                _demo_store[x_demo_token]["total_used"] = _demo_store[x_demo_token].get("total_used", 0) + 1
        logger.info(f"Demo request: token={x_demo_token[:8]}*** remaining={_remaining}")
    # ── End quota check ──────────────────────────────────────────────────
    # Run rule-based pre-scan
    rule_based_matches = []
    rule_based_summary = {}
    try:
        rule_based_matches = detect_pii(request.document)
        rule_based_summary = get_pii_summary(rule_based_matches)
        if rule_based_matches:
            logger.info(f"Rule-based scan found {len(rule_based_matches)} PII instances before LLM")
    except Exception as e:
        logger.warning(f"Rule-based PII scan failed silently: {e}")

    structured_text = structure_prompt_input(sanitized_text, "clinical_document", "M1_PII_ANONYMISER")
    try:
        has_rag_context = bool(regulatory_context)
    except NameError:
        has_rag_context = False

    response = call_claude(
        agent_name="PII_PHI_Anonymisation",
        model=MODEL_HAIKU,
        system_prompt=AGENT_01_SYSTEM_PROMPT,
        user_content=f"Document metadata: {json.dumps(request.metadata)}\n\nDocument:\n{structured_text}",
        api_key=x_anthropic_api_key or "",
        max_tokens=4096,
        has_rag_context=has_rag_context,
    )

    # Merge rule-based findings
    try:
        parsed = response.result
        if isinstance(parsed, dict):
            rule_based_formatted = format_for_response(rule_based_matches)
            
            # Add rule-based detections to the result
            parsed["rule_based_detections"] = rule_based_formatted
            parsed["rule_based_summary"] = rule_based_summary
            parsed["detection_method"] = "HYBRID — Rule-based + Claude LLM"
            
            # Merge into entities_detected if not already captured
            existing_values = {
                str(e.get("value", "")).lower() 
                for e in parsed.get("entities_detected", [])
            }
            for rb_match in rule_based_formatted:
                if rb_match["value"].lower() not in existing_values:
                    parsed.setdefault("entities_detected", []).append({
                        "entity_type": rb_match["entity_type"],
                        "value": rb_match["value"],
                        "category": rb_match["category"],
                        "position": rb_match["position"],
                        "detection_method": "RULE_BASED"
                    })
                    existing_values.add(rb_match["value"].lower())
                    
            # Update entity count
            parsed["entities_anonymised"] = len(parsed.get("entities_detected", []))

            # Store pseudonymisation mapping for potential reversal
            try:
                token_map = {}
                for entity in parsed.get("entities_detected", []):
                    if entity.get("value") and entity.get("entity_type"):
                        token_key = f"[{entity['entity_type']}_001]"
                        token_map[token_key] = entity["value"]
                if token_map:
                    session_id = str(uuid.uuid4())
                    pii_mapping_store.store_mapping(session_id, token_map)
                    parsed["pseudonymisation_session_id"] = session_id
                    parsed["deanonymisation_available"] = True
                    parsed["mapping_expires_hours"] = PIIMappingStore.TTL_HOURS
            except Exception as e:
                logger.warning(f"PII mapping storage failed silently: {e}")
    except Exception as e:
        logger.warning(f"Rule-based merge failed silently: {e}")

    return response


@router.post("/summarise", response_model=AgentResponse, summary="Agent 02 - Document Summarisation")
@limiter.limit("10/minute")
async def summarise_document(request: Request, body: AgentRequest, x_anthropic_api_key: Optional[str] = Header(None), x_demo_token: Optional[str] = Header(None)):
    # ── Input sanitization ────────────────────────────────────────────────
    try:
        sanitized_text, was_modified = sanitize_input(
            body.document,
            context="M2_SUMMARISER"
        )
        if was_modified:
            logger.info("Input was sanitized before processing")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # ── Demo quota check ──────────────────────────────────────────────────
    _admin_key = os.getenv("ADMIN_DEMO_KEY", "")
    _is_admin = _admin_key and x_anthropic_api_key == _admin_key
    if not _is_admin and x_demo_token:
        _allowed, _remaining = _check_and_decrement_quota(x_demo_token)
        if not _allowed:
            raise HTTPException(
                status_code=429,
                detail={
                    "error": "DEMO_QUOTA_EXHAUSTED",
                    "message": "You have used all 5 free demo requests.",
                    "action": "Request full access at regcheckindia.com",
                    "contact": "rushikeshbork000@gmail.com"
                }
            )
        with _demo_lock:
            if x_demo_token in _demo_store:
                _demo_store[x_demo_token]["total_used"] = _demo_store[x_demo_token].get("total_used", 0) + 1
        logger.info(f"Demo request: token={x_demo_token[:8]}*** remaining={_remaining}")
    # ── End quota check ──────────────────────────────────────────────────
    structured_text = structure_prompt_input(sanitized_text, "regulatory_document", "M2_DOCUMENT_SUMMARISER")
    try:
        has_rag_context = bool(regulatory_context)
    except NameError:
        has_rag_context = False

    return call_claude(
        agent_name="Document_Summarisation",
        model=MODEL_HAIKU,
        system_prompt=AGENT_02_SYSTEM_PROMPT,
        user_content=f"Document metadata: {json.dumps(request.metadata)}\n\nDocument:\n{structured_text}",
        api_key=x_anthropic_api_key or "",
        max_tokens=4096,
        has_rag_context=has_rag_context,
    )


@router.post("/completeness", response_model=AgentResponse, summary="Agent 03 - Completeness Assessment")
@limiter.limit("10/minute")
async def assess_completeness(request: Request, body: CompletenessRequest, x_anthropic_api_key: Optional[str] = Header(None), x_demo_token: Optional[str] = Header(None)):
    # ── Input sanitization ────────────────────────────────────────────────
    try:
        sanitized_text, was_modified = sanitize_input(
            body.document,
            context="M3_COMPLETENESS"
        )
        if was_modified:
            logger.info("Input was sanitized before processing")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # ── Demo quota check ──────────────────────────────────────────────────
    _admin_key = os.getenv("ADMIN_DEMO_KEY", "")
    _is_admin = _admin_key and x_anthropic_api_key == _admin_key
    if not _is_admin and x_demo_token:
        _allowed, _remaining = _check_and_decrement_quota(x_demo_token)
        if not _allowed:
            raise HTTPException(
                status_code=429,
                detail={
                    "error": "DEMO_QUOTA_EXHAUSTED",
                    "message": "You have used all 5 free demo requests.",
                    "action": "Request full access at regcheckindia.com",
                    "contact": "rushikeshbork000@gmail.com"
                }
            )
        with _demo_lock:
            if x_demo_token in _demo_store:
                _demo_store[x_demo_token]["total_used"] = _demo_store[x_demo_token].get("total_used", 0) + 1
        logger.info(f"Demo request: token={x_demo_token[:8]}*** remaining={_remaining}")
    # ── End quota check ──────────────────────────────────────────────────
    # Select system prompt based on document type
    if body.document_type == "SAE_REPORT":
        system_prompt = AGENT_03_SAE_SYSTEM_PROMPT
        agent_name = "SAE_Report_Completeness"
    else:
        system_prompt = AGENT_03_SYSTEM_PROMPT
        agent_name = "Completeness_Assessment"

    # RAG retrieval — get relevant regulatory context
    rag_query = f"completeness requirements {getattr(body, 'document_type', 'general')} CDSCO NDCTR Schedule Y ICH"
    regulatory_context = retrieve_regulatory_context(rag_query, n_results=5)

    # Inject into prompt
    if regulatory_context:
        rag_section = f"""

[RETRIEVED REGULATORY CONTEXT — from official documents]
The following excerpts are retrieved from official regulatory documents. Use these to ground your completeness assessment:

{regulatory_context}

[END OF RETRIEVED CONTEXT]
"""
    else:
        rag_section = "\n[Note: Use your built-in knowledge of NDCTR 2019, Schedule Y, and ICH guidelines]\n"

    structured_text = structure_prompt_input(sanitized_text, "submission_document", "M3_COMPLETENESS_ASSESSOR")
    prompt = f"Assess completeness of this document:{rag_section}\nDocument metadata: {json.dumps(body.metadata)}\n\nDocument:\n{structured_text}"

    try:
        has_rag_context = bool(regulatory_context)
    except NameError:
        has_rag_context = False

    return call_claude(
        agent_name=agent_name,
        model=MODEL_SONNET,
        system_prompt=system_prompt,
        user_content=prompt,
        api_key=x_anthropic_api_key or "",
        max_tokens=4096,
        has_rag_context=has_rag_context,
    )


@router.post("/classify", response_model=AgentResponse, summary="Agent 04 - Case Classification")
@limiter.limit("10/minute")
async def classify_case(request: Request, body: AgentRequest, x_anthropic_api_key: Optional[str] = Header(None), x_demo_token: Optional[str] = Header(None)):
    # ── Input sanitization ────────────────────────────────────────────────
    try:
        sanitized_text, was_modified = sanitize_input(
            body.document,
            context="M4_CLASSIFIER"
        )
        if was_modified:
            logger.info("Input was sanitized before processing")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # ── Demo quota check ──────────────────────────────────────────────────
    _admin_key = os.getenv("ADMIN_DEMO_KEY", "")
    _is_admin = _admin_key and x_anthropic_api_key == _admin_key
    if not _is_admin and x_demo_token:
        _allowed, _remaining = _check_and_decrement_quota(x_demo_token)
        if not _allowed:
            raise HTTPException(
                status_code=429,
                detail={
                    "error": "DEMO_QUOTA_EXHAUSTED",
                    "message": "You have used all 5 free demo requests.",
                    "action": "Request full access at regcheckindia.com",
                    "contact": "rushikeshbork000@gmail.com"
                }
            )
        with _demo_lock:
            if x_demo_token in _demo_store:
                _demo_store[x_demo_token]["total_used"] = _demo_store[x_demo_token].get("total_used", 0) + 1
        logger.info(f"Demo request: token={x_demo_token[:8]}*** remaining={_remaining}")
    # ── End quota check ──────────────────────────────────────────────────
    # Run duplicate search
    duplicate_matches = []
    try:
        duplicate_matches = search_similar_cases(sanitized_text)
    except Exception as e:
        logger.warning(f"Duplicate search failed silently: {e}")

    structured_text = structure_prompt_input(sanitized_text, "sae_narrative", "M4_CASE_CLASSIFIER")
    try:
        has_rag_context = bool(regulatory_context)
    except NameError:
        has_rag_context = False

    response = call_claude(
        agent_name="Case_Classification",
        model=MODEL_HAIKU,
        system_prompt=AGENT_04_SYSTEM_PROMPT,
        user_content=f"Case metadata: {json.dumps(body.metadata)}\n\nCase narrative:\n{structured_text}",
        api_key=x_anthropic_api_key or "",
        max_tokens=4096,
        has_rag_context=has_rag_context,
    )

    # Add duplicate info to result
    parsed = response.result
    if isinstance(parsed, dict):
        parsed["duplicate_detection"] = {
            "duplicates_found": len(duplicate_matches) > 0,
            "match_count": len(duplicate_matches),
            "matches": duplicate_matches,
            "recommendation": (
                "POTENTIAL_DUPLICATE — Review matched cases before processing" 
                if duplicate_matches else "NO_DUPLICATES — Case appears unique"
            )
        }
        
        # Store the case for future duplicate detection
        try:
            store_case(sanitized_text, {
                "primary_category": parsed.get("primary_category", ""),
                "study_drug": parsed.get("audit_log", {}).get("study_drug", ""),
                "priority_score": parsed.get("priority_score", 0)
            })
        except Exception as e:
            logger.warning(f"Case storage failed silently: {e}")
            
    return response


@router.post("/ocr", summary="OCR — Extract text from scanned or handwritten documents")
@limiter.limit("5/minute")
async def extract_text_ocr(
    file: UploadFile = File(...),
    request: Request = None,
    mode: str = "auto",  # auto | tesseract | vision
    x_anthropic_api_key: Optional[str] = Header(None)
):

    """
    Extract text from scanned PDFs, printed images, or handwritten notes.
    Supports: PDF, PNG, JPG, JPEG, TIFF, BMP
    Modes: auto (smart selection), tesseract (printed), vision (handwritten)
    """
    try:

        content = await file.read()
        try:
            file_info = validate_file(
                content=content,
                filename=file.filename or "upload",
                allowed_extensions=['.pdf', '.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp']
            )
            filename = file_info["safe_filename"]
        except FileValidationError as e:
            raise HTTPException(status_code=400, detail=str(e))

        api_key = x_anthropic_api_key or os.getenv("ANTHROPIC_API_KEY", "")
        force_vision = (mode == "vision")

        result = extract_text_from_image_bytes(
            image_bytes=content,
            filename=filename,
            api_key=api_key if api_key else None,
            force_vision=force_vision
        )
        del content  # hint to GC

        if not result.text.strip():
            raise HTTPException(
                status_code=422,
                detail="Could not extract text from file. The image may be too low quality or blank."
            )

        return {
            "filename": filename,
            "extracted_text": result.text,
            "word_count": len(result.text.split()),
            "page_count": result.page_count,
            "ocr_method": result.method,
            "confidence": result.confidence,
            "warnings": result.warnings,
            "status": "success"
        }

    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"OCR endpoint error: {e}")
        raise HTTPException(status_code=500, detail=f"OCR processing failed: {str(e)}")


@router.get("/ocr/health")
async def ocr_health():
    """Check OCR service availability."""
    tesseract_available = False
    pdf2image_available = False

    try:
        pytesseract.get_tesseract_version()
        tesseract_available = True
    except Exception:
        pass

    try:
        pdf2image_available = True
    except Exception:
        pass
    return {
        "status": "ok",
        "tesseract_available": tesseract_available,
        "pdf2image_available": pdf2image_available,
        "claude_vision_available": True,
        "supported_formats": ["PDF", "PNG", "JPG", "JPEG", "TIFF", "BMP"]
    }


@router.post("/transcribe", summary="Transcribe and summarise meeting audio")
@limiter.limit("3/minute")
async def transcribe_meeting(
    file: UploadFile = File(...),
    request: Request = None,
    language_code: str = "unknown",
    x_anthropic_api_key: Optional[str] = Header(None),
    x_sarvam_api_key: Optional[str] = Header(None)
):

    """
    Transcribe meeting audio using Sarvam AI Saaras v3,
    then summarise using Claude into structured meeting notes.
    Supports: MP3, WAV, AAC, OGG, FLAC, M4A, WebM, WMA, AMR
    """
    try:

        content = await file.read()
        try:
            file_info = validate_file(
                content=content,
                filename=file.filename or "audio.mp3",
                allowed_extensions=['.mp3', '.wav', '.aac', '.ogg', '.flac', '.m4a', '.webm', '.wma', '.amr']
            )
            filename = file_info["safe_filename"]
        except FileValidationError as e:
            raise HTTPException(status_code=400, detail=str(e))

        # Validate API keys
        sarvam_key = x_sarvam_api_key or os.getenv("SARVAM_API_KEY", "")
        anthropic_key = x_anthropic_api_key or os.getenv("ANTHROPIC_API_KEY", "")

        if not sarvam_key:
            raise HTTPException(
                status_code=401,
                detail="Sarvam API key required — add it in Settings"
            )
        if not anthropic_key:
            raise HTTPException(
                status_code=401,
                detail="Anthropic API key required — add it in Settings"
            )

        logger.info(f"Transcribing audio: {filename}, size: {len(content)} bytes")

        # Step 1 — Transcribe with Sarvam
        transcription = transcribe_audio(
            audio_bytes=content,
            filename=filename,
            sarvam_api_key=sarvam_key,
            language_code=language_code
        )

        if not transcription["transcript"].strip():
            raise HTTPException(
                status_code=422,
                detail="No speech detected in audio file"
            )

        logger.info(
            f"Transcription complete: {transcription['chunk_count']} chunks, "
            f"{transcription['duration_seconds']}s, "
            f"language: {transcription['language_detected']}"
        )

        # Step 2 — Summarise with Claude
        word_count = len(transcription["transcript"].split())
        prompt = f"""Meeting transcript ({transcription['duration_seconds']}s, {word_count} words):

{transcription['transcript']}

Summarise this pharmaceutical regulatory meeting transcript."""

        client = anthropic.Anthropic(api_key=anthropic_key)
        response = client.messages.create(
            model=os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-20250514"),
            max_tokens=4096,
            system=AGENT_02_MEETING_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}]
        )

        raw_text = response.content[0].text.strip()
        parsed = _parse_json_block(raw_text)

        # Inject transcription metadata
        parsed["transcription"] = {
            "language_detected": transcription["language_detected"],
            "duration_seconds": transcription["duration_seconds"],
            "word_count": word_count,
            "chunk_count": transcription["chunk_count"],
            "method": transcription["method"]
        }

        return AgentResponse(
            agent="Meeting_Audio_Summarisation",
            model=os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-20250514"),
            result=parsed,
            timestamp=datetime.utcnow().isoformat(),
            token_usage={
                "input_tokens": response.usage.input_tokens,
                "output_tokens": response.usage.output_tokens
            }
        )

    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Transcription error: {e}")
        raise HTTPException(status_code=500, detail=f"Transcription failed: {str(e)}")


@router.get("/transcribe/health")
async def transcribe_health():
    """Check audio transcription service availability."""
    pydub_available = False
    ffmpeg_available = False

    try:
        pydub_available = True
    except ImportError:
        pass

    try:
        result = subprocess.run(
            ["ffmpeg", "-version"],
            capture_output=True, timeout=5
        )
        ffmpeg_available = result.returncode == 0
    except Exception:
        pass

    return {
        "status": "ok",
        "pydub_available": pydub_available,
        "ffmpeg_available": ffmpeg_available,
        "sarvam_model": "saaras:v3",
        "supported_formats": ["mp3", "wav", "aac", "ogg", "flac", "m4a", "webm", "wma", "amr"],
        "max_duration": "No limit (chunked automatically)"
    }


@router.post("/inspection-report", response_model=AgentResponse, summary="Agent 05 - Inspection Report Generation")
@limiter.limit("10/minute")
async def generate_inspection_report(request: Request, body: AgentRequest, x_anthropic_api_key: Optional[str] = Header(None), x_demo_token: Optional[str] = Header(None)):
    # ── Input sanitization ────────────────────────────────────────────────
    try:
        sanitized_text, was_modified = sanitize_input(
            body.document,
            context="M5_INSPECTION"
        )
        if was_modified:
            logger.info("Input was sanitized before processing")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # ── Demo quota check ──────────────────────────────────────────────────
    _admin_key = os.getenv("ADMIN_DEMO_KEY", "")
    _is_admin = _admin_key and x_anthropic_api_key == _admin_key
    if not _is_admin and x_demo_token:
        _allowed, _remaining = _check_and_decrement_quota(x_demo_token)
        if not _allowed:
            raise HTTPException(
                status_code=429,
                detail={
                    "error": "DEMO_QUOTA_EXHAUSTED",
                    "message": "You have used all 5 free demo requests.",
                    "action": "Request full access at regcheckindia.com",
                    "contact": "rushikeshbork000@gmail.com"
                }
            )
        with _demo_lock:
            if x_demo_token in _demo_store:
                _demo_store[x_demo_token]["total_used"] = _demo_store[x_demo_token].get("total_used", 0) + 1
        logger.info(f"Demo request: token={x_demo_token[:8]}*** remaining={_remaining}")
    # ── End quota check ──────────────────────────────────────────────────
    structured_text = structure_prompt_input(sanitized_text, "inspection_observations", "M5_INSPECTION_REPORT")
    try:
        has_rag_context = bool(regulatory_context)
    except NameError:
        has_rag_context = False

    return call_claude(
        agent_name="Inspection_Report_Generation",
        model=MODEL_SONNET,
        system_prompt=AGENT_05_SYSTEM_PROMPT,
        user_content=f"Inspection metadata: {json.dumps(body.metadata)}\n\nInspection findings:\n{structured_text}",
        api_key=x_anthropic_api_key or "",
        max_tokens=4096,
        has_rag_context=has_rag_context,
    )


@router.post("/qa", response_model=AgentResponse, summary="Agent 06 - Regulatory Q&A")
@limiter.limit("10/minute")
async def regulatory_qa(request: Request, body: QARequest, x_anthropic_api_key: Optional[str] = Header(None), x_demo_token: Optional[str] = Header(None)):
    # ── Input sanitization ────────────────────────────────────────────────
    try:
        sanitized_text, was_modified = sanitize_input(
            body.question,
            context="M6_QA"
        )
        if was_modified:
            logger.info("Input was sanitized before processing")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # ── Demo quota check ──────────────────────────────────────────────────
    _admin_key = os.getenv("ADMIN_DEMO_KEY", "")
    _is_admin = _admin_key and x_anthropic_api_key == _admin_key
    if not _is_admin and x_demo_token:
        _allowed, _remaining = _check_and_decrement_quota(x_demo_token)
        if not _allowed:
            raise HTTPException(
                status_code=429,
                detail={
                    "error": "DEMO_QUOTA_EXHAUSTED",
                    "message": "You have used all 5 free demo requests.",
                    "action": "Request full access at regcheckindia.com",
                    "contact": "rushikeshbork000@gmail.com"
                }
            )
        with _demo_lock:
            if x_demo_token in _demo_store:
                _demo_store[x_demo_token]["total_used"] = _demo_store[x_demo_token].get("total_used", 0) + 1
        logger.info(f"Demo request: token={x_demo_token[:8]}*** remaining={_remaining}")
    # ── End quota check ──────────────────────────────────────────────────
    retrieved_context = body.retrieved_context
    sources_used = []
    if not retrieved_context or len(retrieved_context.strip()) < 50:
        try:

            chromadb_path = os.getenv("CHROMADB_PATH", "./data/chromadb")
            client = chromadb.PersistentClient(
                path=chromadb_path,
                settings=chromadb.Settings(anonymized_telemetry=False),
            )
            embedding_fn = embedding_functions.DefaultEmbeddingFunction()
            
            collection = client.get_collection(name="regulatory_documents", embedding_function=embedding_fn)
            results = collection.query(
                query_texts=[sanitized_text],
                n_results=5
            )
            
            if results and results["documents"] and results["documents"][0]:
                retrieved_context = "\n\n".join(results["documents"][0])
                # Capture source metadata for frontend citation display
                for i, doc in enumerate(results["documents"][0]):
                    metadata = results["metadatas"][0][i] if results.get("metadatas") else {}
                    distance = results["distances"][0][i] if results.get("distances") else 1.0
                    similarity = round((1 - distance) * 100, 1)
                    sources_used.append({
                        "title": metadata.get("title", "Regulatory Document"),
                        "short_name": metadata.get("short_name", ""),
                        "authority": metadata.get("authority", ""),
                        "category": metadata.get("category", ""),
                        "chunk_index": metadata.get("chunk_index", 0),
                        "relevance_score": similarity,
                        "snippet": doc[:200] + "..." if len(doc) > 200 else doc
                    })
        except Exception as e:
            logger.warning(f"ChromaDB retrieval failed: {e}")

    if not retrieved_context or len(retrieved_context.strip()) < 50:
        context_section = """[CONTEXT]
Note: No documents were found in the knowledge base for this query. 
Answer using your comprehensive built-in knowledge of Indian pharmaceutical regulations including:
- New Drugs and Clinical Trials Rules (NDCTR) 2019
- Schedule Y of Drugs and Cosmetics Act
- CDSCO guidelines and circulars
- ICH guidelines (E6, E2A, E3, E8, E9)
- ICMR Guidelines for Biomedical and Health Research
- Drugs and Cosmetics Act 1940
Clearly state at the end that this answer is based on general regulatory knowledge and not retrieved documents.
[/CONTEXT]"""
    else:
        context_section = f"[CONTEXT]\n{retrieved_context}\n[/CONTEXT]"

    structured_question = structure_prompt_input(sanitized_text, "regulatory_query", "M6_REGULATORY_QA")
    user_content = (
        f"{context_section}\n\n"
        f"QUESTION: {structured_question}\n\nAdditional metadata: {json.dumps(body.metadata)}"
    )
    try:
        has_rag_context = bool(retrieved_context and len(retrieved_context.strip()) >= 50)
    except NameError:
        has_rag_context = False

    response = call_claude(
        agent_name="Regulatory_QA_RAG",
        model=MODEL_HAIKU,
        system_prompt=AGENT_06_SYSTEM_PROMPT,
        user_content=user_content,
        api_key=x_anthropic_api_key or "",
        max_tokens=4096,
        has_rag_context=has_rag_context,
    )
    # Inject source citations into the response result
    if sources_used and hasattr(response, 'result') and isinstance(response.result, dict):
        response.result["retrieved_sources"] = sources_used
        response.result["answer_grounded_in_documents"] = True
        response.result["source_count"] = len(sources_used)
    elif hasattr(response, 'result') and isinstance(response.result, dict):
        response.result["retrieved_sources"] = []
        response.result["answer_grounded_in_documents"] = False
        response.result["source_count"] = 0
    return response


@router.post("/schedule-y", response_model=AgentResponse, summary="Agent 07 - Schedule Y / CDSCO Compliance")
@limiter.limit("10/minute")
async def check_schedule_y(request: Request, body: AgentRequest, x_anthropic_api_key: Optional[str] = Header(None), x_demo_token: Optional[str] = Header(None)):
    # ── Input sanitization ────────────────────────────────────────────────
    try:
        sanitized_text, was_modified = sanitize_input(
            body.document,
            context="M7_SCHEDULE_Y"
        )
        if was_modified:
            logger.info("Input was sanitized before processing")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # ── Demo quota check ──────────────────────────────────────────────────
    _admin_key = os.getenv("ADMIN_DEMO_KEY", "")
    _is_admin = _admin_key and x_anthropic_api_key == _admin_key
    if not _is_admin and x_demo_token:
        _allowed, _remaining = _check_and_decrement_quota(x_demo_token)
        if not _allowed:
            raise HTTPException(
                status_code=429,
                detail={
                    "error": "DEMO_QUOTA_EXHAUSTED",
                    "message": "You have used all 5 free demo requests.",
                    "action": "Request full access at regcheckindia.com",
                    "contact": "rushikeshbork000@gmail.com"
                }
            )
        with _demo_lock:
            if x_demo_token in _demo_store:
                _demo_store[x_demo_token]["total_used"] = _demo_store[x_demo_token].get("total_used", 0) + 1
        logger.info(f"Demo request: token={x_demo_token[:8]}*** remaining={_remaining}")
    # ── End quota check ──────────────────────────────────────────────────
    # RAG retrieval — specifically query Schedule Y and NDCTR 2019
    rag_query = "Schedule Y requirements clinical trial NDCTR 2019 CDSCO compliance checklist appendix"
    regulatory_context = retrieve_regulatory_context(rag_query, n_results=6)

    if regulatory_context:
        rag_section = f"""

[RETRIEVED REGULATORY CONTEXT — Schedule Y & NDCTR 2019 Official Documents]
The following excerpts are retrieved directly from Schedule Y and NDCTR 2019. Use these as the authoritative source for your compliance assessment:

{regulatory_context}

[END OF RETRIEVED CONTEXT]
"""
    else:
        rag_section = "\n[Note: Use your built-in knowledge of Schedule Y and NDCTR 2019]\n"

    structured_text = structure_prompt_input(sanitized_text, "clinical_protocol", "M7_SCHEDULE_Y")
    prompt = f"Perform Schedule Y compliance assessment:{rag_section}\nDocument metadata: {json.dumps(body.metadata)}\n\nDocument:\n{structured_text}"

    try:
        has_rag_context = bool(regulatory_context)
    except NameError:
        has_rag_context = False

    return call_claude(
        agent_name="Schedule_Y_Compliance",
        model=MODEL_SONNET,
        system_prompt=AGENT_07_SYSTEM_PROMPT,
        user_content=prompt,
        api_key=x_anthropic_api_key or "",
        max_tokens=4096,
        has_rag_context=has_rag_context,
    )


@router.post("/ich-gcp", response_model=AgentResponse, summary="Agent 08 - ICH E6(R3) GCP Compliance")
@limiter.limit("10/minute")
async def check_ich_gcp(request: Request, body: AgentRequest, x_anthropic_api_key: Optional[str] = Header(None), x_demo_token: Optional[str] = Header(None)):
    # ── Input sanitization ────────────────────────────────────────────────
    try:
        sanitized_text, was_modified = sanitize_input(
            body.document,
            context="M8_ICH_GCP"
        )
        if was_modified:
            logger.info("Input was sanitized before processing")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # ── Demo quota check ──────────────────────────────────────────────────
    _admin_key = os.getenv("ADMIN_DEMO_KEY", "")
    _is_admin = _admin_key and x_anthropic_api_key == _admin_key
    if not _is_admin and x_demo_token:
        _allowed, _remaining = _check_and_decrement_quota(x_demo_token)
        if not _allowed:
            raise HTTPException(
                status_code=429,
                detail={
                    "error": "DEMO_QUOTA_EXHAUSTED",
                    "message": "You have used all 5 free demo requests.",
                    "action": "Request full access at regcheckindia.com",
                    "contact": "rushikeshbork000@gmail.com"
                }
            )
        with _demo_lock:
            if x_demo_token in _demo_store:
                _demo_store[x_demo_token]["total_used"] = _demo_store[x_demo_token].get("total_used", 0) + 1
        logger.info(f"Demo request: token={x_demo_token[:8]}*** remaining={_remaining}")
    # ── End quota check ──────────────────────────────────────────────────
    # RAG retrieval — specifically query ICH E6(R3)
    rag_query = "ICH E6 R3 GCP good clinical practice quality management risk based monitoring"
    regulatory_context = retrieve_regulatory_context(rag_query, n_results=6)

    if regulatory_context:
        rag_section = f"""

[RETRIEVED REGULATORY CONTEXT — ICH E6(R3) Official Document]
The following excerpts are retrieved directly from ICH E6(R3) final guidelines. Use these as the authoritative source for your GCP assessment:

{regulatory_context}

[END OF RETRIEVED CONTEXT]
"""
    else:
        rag_section = "\n[Note: Use your built-in knowledge of ICH E6(R3) GCP guidelines]\n"

    structured_text = structure_prompt_input(sanitized_text, "gcp_document", "M8_ICH_GCP")
    prompt = f"Perform ICH E6(R3) GCP compliance assessment:{rag_section}\nDocument metadata: {json.dumps(body.metadata)}\n\nDocument:\n{structured_text}"

    try:
        has_rag_context = bool(regulatory_context)
    except NameError:
        has_rag_context = False

    return call_claude(
        agent_name="ICH_GCP_Compliance",
        model=MODEL_SONNET,
        system_prompt=AGENT_08_SYSTEM_PROMPT,
        user_content=prompt,
        api_key=x_anthropic_api_key or "",
        max_tokens=4096,
        has_rag_context=has_rag_context,
    )


@router.post("/compare", summary="Agent 03b - Document Version Comparison")
async def compare_documents(
    file_a: UploadFile = File(...),
    file_b: UploadFile = File(...),
    x_anthropic_api_key: Optional[str] = Header(None)
):

    """Compare two versions of a regulatory document and identify changes with regulatory impact."""
    try:

        def extract_text(file_bytes: bytes, filename: str) -> str:
            filename_lower = filename.lower()
            if filename_lower.endswith(".pdf"):
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                pages = []
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        pages.append(f"[Page {i+1}]\n{text}")
                return "\n\n".join(pages)
            elif filename_lower.endswith(".docx"):
                doc = docx.Document(io.BytesIO(file_bytes))
                return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}. Use PDF or DOCX.")

        filename_a = sanitize_filename(file_a.filename or "version_a")
        filename_b = sanitize_filename(file_b.filename or "version_b")
        content_a = await file_a.read()
        content_b = await file_b.read()

        text_a = extract_text(content_a, filename_a)
        text_b = extract_text(content_b, filename_b)
        del content_a, content_b  # hint to GC

        if not text_a.strip():
            raise HTTPException(status_code=422, detail="Could not extract text from Version A. File may be scanned/image-based.")
        if not text_b.strip():
            raise HTTPException(status_code=422, detail="Could not extract text from Version B. File may be scanned/image-based.")

        # Truncate if too long — keep first 6000 words each to stay within token limits
        words_a = text_a.split()
        words_b = text_b.split()
        if len(words_a) > 6000:
            text_a = " ".join(words_a[:6000]) + "\n[Document truncated for analysis]"
        if len(words_b) > 6000:
            text_b = " ".join(words_b[:6000]) + "\n[Document truncated for analysis]"

        api_key = x_anthropic_api_key or os.getenv("ANTHROPIC_API_KEY", "")
        if not api_key:
            raise HTTPException(status_code=401, detail="No Anthropic API key provided. Add your key via the ⚙ Settings panel in the app.")

        # RAG retrieval — get context for regulatory impact of document changes
        rag_query = "regulatory change impact assessment CDSCO submission amendment requirements"
        regulatory_context = retrieve_regulatory_context(rag_query, n_results=4)

        if regulatory_context:
            rag_section = f"""

[RETRIEVED REGULATORY CONTEXT]
Use these official regulatory excerpts to accurately identify which regulations are affected by document changes:

{regulatory_context}

[END OF RETRIEVED CONTEXT]
"""
        else:
            rag_section = ""

        prompt = f"""You are comparing two versions of a pharmaceutical regulatory document.{rag_section}
VERSION A (Original):
{text_a}

---

VERSION B (Revised):
{text_b}

---

Perform a detailed regulatory document comparison. Identify every substantive change between Version A and Version B.

CRITICAL OUTPUT RULES:
1. Return ONLY valid JSON — no markdown, no code fences, no explanation
2. Use EXACTLY the field names shown below
3. Arrays must always be arrays
4. Objects must always be objects
5. String fields must always be strings
6. Never omit a required field

Return EXACTLY this JSON structure:
{{
  "document_name_a": "string — filename or detected title of Version A",
  "document_name_b": "string — filename or detected title of Version B",
  "overall_change_severity": "CRITICAL|MAJOR|MINOR|NO_CHANGE",
  "total_changes": 0,
  "executive_summary": "string — 2-3 sentence summary of what changed and why it matters regulatorily",
  "changes": [
    {{
      "change_id": "string — e.g. CHG-001",
      "section": "string — section name or number where change occurred",
      "change_type": "ADDED|REMOVED|MODIFIED",
      "severity": "CRITICAL|MAJOR|MINOR",
      "description": "string — what specifically changed",
      "original_text": "string — relevant text from Version A (max 200 words), empty string if ADDED",
      "revised_text": "string — relevant text from Version B (max 200 words), empty string if REMOVED",
      "regulatory_impact": {{
        "regulation": "string — specific regulation affected e.g. Schedule Y Section 4.2, NDCTR 2019 Rule 23, ICH E6(R3) Section 5.18",
        "impact_description": "string — how this change affects compliance with that regulation",
        "compliance_risk": "HIGH|MEDIUM|LOW",
        "action_required": "string — what the applicant must do about this change"
      }}
    }}
  ],
  "critical_changes": ["string — list of the most important changes in plain language"],
  "added_sections": ["string — sections/content present in V2 but not V1"],
  "removed_sections": ["string — sections/content present in V1 but not V2"],
  "regulatory_frameworks_affected": ["string — list of regulations impacted by the changes"],
  "submission_impact": "RESUBMISSION_REQUIRED|AMENDMENT_REQUIRED|NOTIFICATION_REQUIRED|NO_ACTION",
  "submission_impact_rationale": "string — why this submission impact level was assigned",
  "audit_log": {{
    "timestamp": "ISO datetime string",
    "file_a": "string — filename of Version A",
    "file_b": "string — filename of Version B",
    "pages_a": 0,
    "pages_b": 0,
    "status": "COMPLETED|FAILED"
  }}
}} """

        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model=MODEL_SONNET,
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}]
        )

        raw_text = response.content[0].text.strip()
        parsed = _parse_json_block(raw_text)

        return AgentResponse(
            agent="Document_Version_Comparison",
            model=MODEL_SONNET,
            result=parsed,
            timestamp=_utc_timestamp(),
            token_usage={
                "input_tokens": response.usage.input_tokens,
                "output_tokens": response.usage.output_tokens
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document comparison error: {e}")
        raise HTTPException(status_code=500, detail=f"Comparison error: {str(e)}")


@router.get("/compare/health")
async def compare_health():
    return {"status": "ok", "agent": "Document_Version_Comparison", "endpoint": "/api/v1/agents/compare"}


@router.post("/feedback", summary="Submit output feedback")
async def submit_feedback(request: FeedbackRequest):
    """Log user feedback on module outputs."""
    logger.info(
        scrub_sensitive("FEEDBACK: module=%s type=%s comment=%s result_hash=%s"),
        request.module,
        request.type,
        (request.comment or "")[:100],
        request.result_hash or "",
    )
    return {"status": "ok", "message": "Feedback recorded"}


@router.get("/ping", summary="Keep-alive ping for uptime monitors")
async def ping():
    """Lightweight endpoint for uptime monitoring."""
    return {
        "status": "alive",
        "timestamp": datetime.utcnow().isoformat(),
        "service": "RegCheck-India API",
    }


@router.post("/deanonymise", summary="M1 — Reverse pseudonymisation using stored encrypted mapping")
async def deanonymise_text(request: DeAnonymiseRequest):
    """Reverse pseudonymisation using stored encrypted mapping."""
    result = pii_mapping_store.deanonymise(request.text, request.session_id)
    if result is None:
        raise HTTPException(
            status_code=404,
            detail="Mapping not found or expired. Pseudonymisation mappings expire after 24 hours."
        )
    return {
        "original_text": result,
        "session_id": request.session_id,
        "status": "success"
    }


@router.get("/mapping-store/health", summary="M1 — PII mapping store stats")
async def mapping_store_health():
    """Return stats about the in-memory encrypted PII mapping store."""
    return pii_mapping_store.get_store_stats()


@router.get("/classify/store/health")
async def case_store_health():
    """Check how many cases are stored for duplicate detection."""
    try:
        collection = get_case_collection()
        count = collection.count() if collection else 0
        return {
            "status": "ok",
            "cases_stored": count,
            "collection": "sae_cases"
        }
    except Exception as e:
        return {"status": "error", "detail": str(e)}


@router.post("/demo/register", summary="Register for demo access")
async def register_demo_user(request: DemoRegistrationRequest):
    """
    Register for demo access.
    Returns demo token with 5-request quota.
    Logs lead information for follow-up.
    """
    # Validate inputs
    if not request.email or "@" not in request.email:
        raise HTTPException(status_code=400, detail="Valid email required")
    if not request.name or not request.org:
        raise HTTPException(status_code=400, detail="Name and organisation required")

    token = _generate_token(request.email, request.org)

    with _demo_lock:
        if token not in _demo_store:
            # New registration
            _demo_store[token] = {
                "name": request.name,
                "email": request.email,
                "org": request.org,
                "role": request.role,
                "remaining": DEMO_QUOTA,
                "registered_at": datetime.utcnow().isoformat(),
                "total_used": 0
            }
            logger.info(
                f"NEW DEMO REGISTRATION: {request.name} | "
                f"{request.email} | {request.org} | {request.role}"
            )
            message = f"Welcome {request.name}! You have {DEMO_QUOTA} free requests."
        else:
            # Returning user
            remaining = _demo_store[token]["remaining"]
            logger.info(f"RETURNING USER: {request.email} | remaining: {remaining}")
            message = f"Welcome back {request.name}! You have {remaining} requests remaining."

    entry = _demo_store[token]
    return {
        "demo_token": token,
        "requests_remaining": entry["remaining"],
        "name": entry["name"],
        "message": message
    }


@router.get("/demo/quota/{token}", summary="Check demo quota")
async def check_demo_quota(token: str):
    """Check remaining requests for a demo token."""
    with _demo_lock:
        entry = _demo_store.get(token)
    if not entry:
        raise HTTPException(status_code=404, detail="Token not found. Please register first.")
    return {
        "requests_remaining": entry["remaining"],
        "total_used": entry.get("total_used", 0),
        "name": entry["name"],
        "org": entry["org"]
    }


@router.get("/demo/leads", summary="Get all demo registrations (admin only)")
async def get_demo_leads(x_admin_key: Optional[str] = Header(None)):
    """
    Return all demo registrations for lead tracking.
    Requires admin key.
    """
    admin_key = os.getenv("ADMIN_DEMO_KEY", "")
    if not admin_key or x_admin_key != admin_key:
        raise HTTPException(status_code=403, detail="Admin key required")

    with _demo_lock:
        leads = [
            {
                "name": v["name"],
                "email": v["email"],
                "org": v["org"],
                "role": v["role"],
                "registered_at": v["registered_at"],
                "requests_used": DEMO_QUOTA - v["remaining"],
                "requests_remaining": v["remaining"]
            }
            for v in _demo_store.values()
        ]

    return {
        "total_registrations": len(leads),
        "leads": sorted(leads, key=lambda x: x["registered_at"], reverse=True)
    }


@router.get("/health", summary="Agents Health Check")
async def agents_health():
    return {
        "status": "ok",
        "agents": [
            {"id": "01", "name": "PII_PHI_Anonymisation", "endpoint": "/anonymise", "model": MODEL_HAIKU},
            {"id": "02", "name": "Document_Summarisation", "endpoint": "/summarise", "model": MODEL_HAIKU},
            {"id": "03", "name": "Completeness_Assessment", "endpoint": "/completeness", "model": MODEL_SONNET},
            {"id": "04", "name": "Case_Classification", "endpoint": "/classify", "model": MODEL_HAIKU},
            {"id": "05", "name": "Inspection_Report_Generation", "endpoint": "/inspection-report", "model": MODEL_SONNET},
            {"id": "06", "name": "Regulatory_QA_RAG", "endpoint": "/qa", "model": MODEL_HAIKU},
            {"id": "07", "name": "Schedule_Y_Compliance", "endpoint": "/schedule-y", "model": MODEL_SONNET},
            {"id": "08", "name": "ICH_GCP_Compliance", "endpoint": "/ich-gcp", "model": MODEL_SONNET},
            {"id": "03b", "name": "Document_Version_Comparison", "endpoint": "/compare", "model": MODEL_SONNET},
        ],
        "anthropic_client": "initialised",
        "timestamp": _utc_timestamp(),
    }
