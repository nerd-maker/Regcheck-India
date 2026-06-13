import io
import logging
from datetime import datetime, timezone
from typing import Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ── Brand colours ─────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1A, 0x37, 0x6C)   # RegCheck dark blue — headings
ACCENT     = RGBColor(0x2E, 0x75, 0xB6)   # Mid blue — sub-headings, dividers
RED        = RGBColor(0xC0, 0x00, 0x00)   # Critical / high risk
AMBER      = RGBColor(0xED, 0x7D, 0x31)   # Major / medium risk
GREEN      = RGBColor(0x37, 0x86, 0x44)   # Minor / low risk / compliant
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)  # Table header background
MID_GREY   = RGBColor(0x59, 0x56, 0x59)  # Body text secondary

# ── Agent display names ────────────────────────────────────────────────────────
AGENT_DISPLAY_NAMES: dict[str, str] = {
    "Schedule_Y_Compliance":       "Schedule Y Compliance Check",
    "ICH_GCP_Compliance":          "ICH E6(R3) GCP Compliance Check",
    "Document_Completeness":       "Document Completeness Check",
    "PII_PHI_Anonymisation":       "PII / PHI Anonymisation Report",
    "SAE_Case_Classification":     "SAE Case Classification Report",
    "Document_Summarisation":      "Document Summarisation Report",
    "Regulatory_QA":               "Regulatory Q&A Report",
    "Inspection_Report":           "Inspection Readiness Report",
    "Cross_Document_Consistency":  "Cross-Document Consistency Check",
}


# ══════════════════════════════════════════════════════════════════════════════
# PUBLIC ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

def generate_word_report(agent_response: dict) -> bytes:
    """
    Generate a Word (.docx) compliance report from an AgentResponse dict.
    Returns raw bytes of the .docx file.
    """
    agent_name: str = agent_response.get("agent", "Unknown Agent")
    model: str = agent_response.get("model", "")
    result: dict = agent_response.get("result", {})
    timestamp: str = agent_response.get("timestamp", datetime.now(timezone.utc).isoformat())
    token_usage: dict = agent_response.get("token_usage", {})

    doc = Document()
    _configure_page(doc)
    _add_header_banner(doc, agent_name, timestamp, model)
    _add_divider(doc)

    # Dispatch to agent-specific section builder
    agent_lower = agent_name.lower()
    if "schedule_y" in agent_lower or "schedule y" in agent_lower:
        _build_schedule_y_body(doc, result)
    elif "ich" in agent_lower or "gcp" in agent_lower:
        _build_ich_gcp_body(doc, result)
    elif "completeness" in agent_lower:
        _build_completeness_body(doc, result)
    elif "pii" in agent_lower or "anonymis" in agent_lower:
        _build_pii_body(doc, result)
    elif "sae" in agent_lower or "classif" in agent_lower:
        _build_sae_body(doc, result)
    elif "summar" in agent_lower:
        _build_summariser_body(doc, result)
    else:
        _build_generic_body(doc, result)

    _add_footer_disclaimer(doc, token_usage)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT SETUP
# ══════════════════════════════════════════════════════════════════════════════

def _configure_page(doc: Document) -> None:
    """A4 page, 2.5cm margins all sides."""
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Set default font for the document
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _add_header_banner(doc: Document, agent_name: str, timestamp: str, model: str) -> None:
    """Blue banner with report title and metadata."""
    display_name = AGENT_DISPLAY_NAMES.get(agent_name, agent_name.replace("_", " ").title())

    # Report title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run("RegCheck-India")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = ACCENT
    run.font.bold = True

    # Agent name as H1
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h1.add_run(display_name)
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # Metadata line
    try:
        dt = datetime.fromisoformat(timestamp.replace("Z", "+00:00"))
        formatted_ts = dt.strftime("%d %B %Y, %H:%M UTC")
    except Exception:
        formatted_ts = timestamp

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run = meta.add_run(f"Generated: {formatted_ts}    |    Model: {model}    |    Requires RA sign-off before regulatory use")
    meta_run.font.name = "Calibri"
    meta_run.font.size = Pt(8)
    meta_run.font.color.rgb = MID_GREY
    meta_run.font.italic = True


def _add_divider(doc: Document) -> None:
    """Horizontal rule using paragraph bottom border."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def _add_section_heading(doc: Document, text: str, level: int = 2) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper() if level == 2 else text)
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.color.rgb = NAVY if level == 2 else ACCENT
    run.font.size = Pt(11) if level == 2 else Pt(10)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def _add_score_line(doc: Document, label: str, score: float | None, out_of: int = 100) -> None:
    """Render a compliance score with colour coding."""
    if score is None:
        return
    score_int = int(round(score))
    color = GREEN if score_int >= 80 else (AMBER if score_int >= 60 else RED)
    verdict = "COMPLIANT" if score_int >= 80 else ("NEEDS REVISION" if score_int >= 60 else "NON-COMPLIANT")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    label_run = p.add_run(f"{label}: ")
    label_run.font.name = "Calibri"
    label_run.font.bold = True
    label_run.font.size = Pt(12)
    score_run = p.add_run(f"{score_int} / {out_of}  —  {verdict}")
    score_run.font.name = "Calibri"
    score_run.font.bold = True
    score_run.font.size = Pt(12)
    score_run.font.color.rgb = color


def _add_findings_list(doc: Document, items: list[Any], severity_color: RGBColor) -> None:
    """Render a bulleted list of finding strings."""
    if not items:
        p = doc.add_paragraph()
        run = p.add_run("None identified.")
        run.font.color.rgb = GREEN
        run.font.italic = True
        run.font.size = Pt(10)
        return
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        text = str(item) if not isinstance(item, str) else item
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = severity_color


def _add_checklist_table(
    doc: Document,
    rows: list[dict],
    columns: list[tuple[str, str, int]],  # (header_label, dict_key, width_dxa)
) -> None:
    """
    Render a structured checklist as a Word table.
    columns: list of (header_label, result_dict_key, column_width_in_dxa)
    """
    if not rows:
        return

    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, (label, _, width) in enumerate(columns):
        cell = hdr_cells[i]
        cell.width = width
        # Navy background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1A376C")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.font.name = "Calibri"
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for i, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        fill = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        for j, (_, key, width) in enumerate(columns):
            cell = row_cells[j]
            cell.width = width
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
            val = row_data.get(key, "")
            if isinstance(val, bool):
                val = "Yes" if val else "No"
            p = cell.paragraphs[0]
            run = p.add_run(str(val) if val is not None else "—")
            run.font.name = "Calibri"
            run.font.size = Pt(9)


def _add_footer_disclaimer(doc: Document, token_usage: dict) -> None:
    """Footer: AI disclaimer + token usage."""
    _add_divider(doc)
    p = doc.add_paragraph()
    input_tokens  = token_usage.get("input_tokens", "—")
    output_tokens = token_usage.get("output_tokens", "—")
    run = p.add_run(
        f"DISCLAIMER: This report is AI-generated and requires review and sign-off by a qualified Regulatory Affairs professional "
        f"before use in any regulatory submission. RegCheck-India | Tokens used: {input_tokens} in / {output_tokens} out"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = MID_GREY
    run.font.italic = True


# ══════════════════════════════════════════════════════════════════════════════
# AGENT-SPECIFIC BODY BUILDERS
# ══════════════════════════════════════════════════════════════════════════════

def _build_schedule_y_body(doc: Document, result: dict) -> None:
    score = result.get("compliance_score")
    _add_score_line(doc, "Schedule Y Compliance Score", score)

    _add_section_heading(doc, "Critical Non-Compliances")
    _add_findings_list(doc, result.get("critical_non_compliances", []), RED)

    _add_section_heading(doc, "Major Non-Compliances")
    _add_findings_list(doc, result.get("major_non_compliances", []), AMBER)

    _add_section_heading(doc, "Minor Non-Compliances")
    _add_findings_list(doc, result.get("minor_non_compliances", []), GREEN)

    checklist = result.get("compliance_checklist", [])
    if checklist:
        _add_section_heading(doc, "Detailed Compliance Checklist")
        _add_checklist_table(doc, checklist, [
            ("Requirement",        "requirement",       3500),
            ("Section",            "section",           1500),
            ("Status",             "status",            1200),
            ("Finding",            "finding",           2500),
            ("Corrective Action",  "corrective_action", 2500),
        ])


def _build_ich_gcp_body(doc: Document, result: dict) -> None:
    score = result.get("gcp_score")
    _add_score_line(doc, "ICH E6(R3) GCP Score", score)

    _add_section_heading(doc, "Critical Deviations")
    _add_findings_list(doc, result.get("critical_deviations", []), RED)

    _add_section_heading(doc, "Major Deviations")
    _add_findings_list(doc, result.get("major_deviations", []), AMBER)

    _add_section_heading(doc, "Minor Deviations")
    _add_findings_list(doc, result.get("minor_deviations", []), GREEN)

    principles = result.get("gcp_principles", [])
    if principles:
        _add_section_heading(doc, "GCP Principles Assessment")
        _add_checklist_table(doc, principles, [
            ("Principle",          "principle",         3000),
            ("ICH Reference",      "ich_reference",     1500),
            ("Status",             "status",            1200),
            ("Observation",        "observation",       2500),
            ("Corrective Action",  "corrective_action", 2500),
        ])


def _build_completeness_body(doc: Document, result: dict) -> None:
    raw_score = result.get("overall_completeness_score")
    if raw_score is not None:
        f = float(raw_score)
        score = f * 100 if f <= 1.0 else f
    else:
        score = None
    _add_score_line(doc, "Completeness Score", score)

    readiness = result.get("submission_readiness", "")
    if readiness:
        p = doc.add_paragraph()
        run = p.add_run(f"Submission Readiness: {readiness}")
        run.font.name = "Calibri"
        run.font.bold = True
        run.font.size = Pt(11)
        color = GREEN if readiness == "READY" else (AMBER if readiness == "NEEDS_REVISION" else RED)
        run.font.color.rgb = color

    _add_section_heading(doc, "Critical Gaps")
    _add_findings_list(doc, result.get("critical_gaps", []), RED)

    _add_section_heading(doc, "Minor Gaps")
    _add_findings_list(doc, result.get("minor_gaps", []), AMBER)

    _add_section_heading(doc, "Missing Sections")
    _add_findings_list(doc, result.get("missing_sections", []), AMBER)


def _build_pii_body(doc: Document, result: dict) -> None:
    report = result.get("anonymisation_report", {})
    count  = result.get("entities_anonymised", 0)

    # Summary box
    p = doc.add_paragraph()
    run = p.add_run(f"Total Entities Anonymised: {count}")
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.size = Pt(12)
    color = GREEN if count == 0 else AMBER
    run.font.color.rgb = color

    if report.get("summary"):
        _add_section_heading(doc, "Summary")
        p = doc.add_paragraph()
        run = p.add_run(str(report["summary"]))
        run.font.name = "Calibri"
        run.font.size = Pt(10)

    # PII/PHI flags
    p = doc.add_paragraph()
    pii_run = p.add_run(f"PII Removed: {'Yes' if report.get('pii_removed') else 'No'}    |    PHI Removed: {'Yes' if report.get('phi_removed') else 'No'}")
    pii_run.font.name = "Calibri"
    pii_run.font.size = Pt(10)
    pii_run.font.italic = True

    entities = result.get("entities_detected", [])
    if entities:
        _add_section_heading(doc, "Entities Detected")
        _add_checklist_table(doc, entities[:50], [  # cap at 50 for readability
            ("Entity Type", "entity_type", 2000),
            ("Value",       "value",       3000),
            ("Category",    "category",   2000),
            ("Position",    "position",   2000),
        ])
    else:
        _add_section_heading(doc, "Entities Detected")
        p = doc.add_paragraph()
        run = p.add_run("No PII/PHI entities detected in this document.")
        run.font.color.rgb = GREEN
        run.font.italic = True


def _build_sae_body(doc: Document, result: dict) -> None:
    confidence = result.get("confidence")
    if confidence is not None:
        score = float(confidence) * 100
        _add_score_line(doc, "Classification Confidence", score)

    primary = result.get("primary_category", "")
    if primary:
        _add_section_heading(doc, "Primary Classification")
        p = doc.add_paragraph()
        run = p.add_run(primary)
        run.font.name = "Calibri"
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = NAVY

    causality = result.get("causality", {})
    if causality:
        _add_section_heading(doc, "Causality Assessment")
        p = doc.add_paragraph()
        run = p.add_run(str(causality.get("assessment", "")))
        run.font.name = "Calibri"
        run.font.size = Pt(10)

    seriousness = result.get("seriousness_criteria", {})
    if seriousness:
        _add_section_heading(doc, "Seriousness Criteria")
        for criterion, met in seriousness.items():
            p = doc.add_paragraph(style="List Bullet")
            label = criterion.replace("_", " ").title()
            run = p.add_run(f"{label}: {'MET' if met else 'Not met'}")
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = RED if met else GREEN

    _add_section_heading(doc, "Flags")
    _add_findings_list(doc, result.get("flags", []), AMBER)

    _add_section_heading(doc, "Regulatory Actions Required")
    _add_findings_list(doc, result.get("regulatory_actions_required", []), RED)


def _build_summariser_body(doc: Document, result: dict) -> None:
    risk = result.get("risk_level", "")
    if risk:
        risk_color = RED if risk == "HIGH" else (AMBER if risk == "MEDIUM" else GREEN)
        p = doc.add_paragraph()
        run = p.add_run(f"Risk Level: {risk}")
        run.font.name = "Calibri"
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = risk_color

    doc_type = result.get("document_type", "")
    if doc_type:
        p = doc.add_paragraph()
        run = p.add_run(f"Document Type: {doc_type}")
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True

    word_count = result.get("word_count_original")
    if word_count:
        p = doc.add_paragraph()
        run = p.add_run(f"Word Count: {word_count:,}")
        run.font.name = "Calibri"
        run.font.size = Pt(10)

    summary = result.get("summary", "")
    if summary:
        _add_section_heading(doc, "Executive Summary")
        p = doc.add_paragraph()
        run = p.add_run(summary)
        run.font.name = "Calibri"
        run.font.size = Pt(10)

    key_sections = result.get("key_sections", [])
    if key_sections:
        _add_section_heading(doc, "Key Sections")
        _add_findings_list(doc, key_sections, NAVY)

    gaps = result.get("compliance_gaps", [])
    if gaps:
        _add_section_heading(doc, "Compliance Gaps Identified")
        _add_findings_list(doc, gaps, RED)

    recs = result.get("recommendations", [])
    if recs:
        _add_section_heading(doc, "Recommendations")
        _add_findings_list(doc, recs, AMBER)

    refs = result.get("regulatory_references", [])
    if refs:
        _add_section_heading(doc, "Regulatory References")
        _add_findings_list(doc, refs, ACCENT)


def _build_generic_body(doc: Document, result: dict) -> None:
    """Fallback for unknown agent types — render all top-level fields."""
    for key, value in result.items():
        if key.startswith("_"):
            continue
        _add_section_heading(doc, key.replace("_", " ").title())
        if isinstance(value, list):
            _add_findings_list(doc, value, NAVY)
        elif isinstance(value, dict):
            for k, v in value.items():
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{k.replace('_', ' ').title()}: {v}")
                run.font.name = "Calibri"
                run.font.size = Pt(10)
        else:
            p = doc.add_paragraph()
            run = p.add_run(str(value))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
