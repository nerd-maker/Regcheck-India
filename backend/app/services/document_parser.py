"""
Document parser service for extracting text from PDF and DOCX files.
"""
import re
from pathlib import Path
from typing import Dict, List, Tuple
from docx import Document
from pypdf import PdfReader


class DocumentParser:
    """Parse pharmaceutical documents and extract structured content."""
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx']
    
    def parse_file(self, file_path: str) -> Dict[str, any]:
        """
        Parse a document file and extract structured content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing:
                - full_text: Complete document text
                - sections: List of (heading, content) tuples
                - metadata: Extracted metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._parse_pdf(file_path)
        elif extension == '.docx':
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, any]:
        """Extract text from PDF file."""
        full_text = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text)
        
        combined_text = '\n'.join(full_text)
        sections = self._extract_sections(combined_text)
        metadata = self._extract_metadata(combined_text)
        
        return {
            'full_text': combined_text,
            'sections': sections,
            'metadata': metadata,
            'page_count': len(full_text)
        }
    
    def _parse_docx(self, file_path: Path) -> Dict[str, any]:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        
        full_text = []
        sections = []
        current_heading = "Introduction"
        current_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check if paragraph is a heading
            if para.style.name.startswith('Heading'):
                # Save previous section
                if current_content:
                    sections.append((current_heading, '\n'.join(current_content)))
                
                current_heading = text
                current_content = []
            else:
                current_content.append(text)
                full_text.append(text)
        
        # Add last section
        if current_content:
            sections.append((current_heading, '\n'.join(current_content)))
        
        combined_text = '\n'.join(full_text)
        metadata = self._extract_metadata(combined_text)
        
        return {
            'full_text': combined_text,
            'sections': sections,
            'metadata': metadata,
            'paragraph_count': len(doc.paragraphs)
        }
    
    def _extract_sections(self, text: str) -> List[Tuple[str, str]]:
        """
        Extract sections from text based on common heading patterns.
        
        Looks for patterns like:
        - 1. Introduction
        - Section 1: Background
        - BACKGROUND
        """
        sections = []
        
        # Common section heading patterns
        patterns = [
            r'^(\d+\.?\s+[A-Z][^\n]+)',  # 1. Introduction or 1 Introduction
            r'^([A-Z][A-Z\s]+)$',  # ALL CAPS HEADINGS
            r'^(Section\s+\d+[:\.]?\s*[^\n]+)',  # Section 1: Title
        ]
        
        lines = text.split('\n')
        current_heading = "Document Content"
        current_content = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Check if line matches any heading pattern
            is_heading = False
            for pattern in patterns:
                if re.match(pattern, line):
                    # Save previous section
                    if current_content:
                        sections.append((current_heading, '\n'.join(current_content)))
                    
                    current_heading = line
                    current_content = []
                    is_heading = True
                    break
            
            if not is_heading:
                current_content.append(line)
        
        # Add last section
        if current_content:
            sections.append((current_heading, '\n'.join(current_content)))
        
        return sections if sections else [("Document Content", text)]
    
    def _extract_metadata(self, text: str) -> Dict[str, str]:
        """Extract common metadata fields from document text."""
        metadata = {}
        
        # Common metadata patterns
        patterns = {
            'protocol_number': r'Protocol\s+(?:Number|ID|No\.?)[\s:]+([A-Z0-9\-]+)',
            'version': r'Version[\s:]+([0-9\.]+)',
            'date': r'Date[\s:]+(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
            'sponsor': r'Sponsor[\s:]+([^\n]+)',
            'title': r'Title[\s:]+([^\n]+)',
        }
        
        for key, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                metadata[key] = match.group(1).strip()
        
        return metadata


# Global parser instance
document_parser = DocumentParser()
